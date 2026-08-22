"""Generate synthetic unstructured banking documents for V0."""

from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

DOCUMENTS = {
    "fraud_policy_v1": {
        "title": "Meridian Bank Card Fraud Policy v1",
        "body": [
            "Purpose: define minimum controls for identifying and responding to suspected card fraud.",
            "High-risk indicators include unfamiliar devices, unusual geography, rapid repeat spending and failed authentication.",
            "An investigator must preserve transaction, authentication and device evidence before closing a case.",
            "Confirmed fraud must be recorded against the original transaction after customer or network confirmation.",
        ],
    },
    "dispute_procedure_v1": {
        "title": "Meridian Bank Dispute Procedure v1",
        "body": [
            "Capture the disputed transaction identifier, customer statement and merchant descriptor.",
            "Check authorization and authentication evidence before submitting a network dispute.",
            "Escalate missing evidence or conflicting customer statements to Fraud Operations.",
        ],
    },
    "escalation_matrix_v1": {
        "title": "Fraud Operations Escalation Matrix v1",
        "body": [
            "P2: suspected account takeover or aggregate exposure above INR 100,000; notify the duty manager.",
            "P3: confirmed single-card fraud below INR 100,000; assign to the card fraud queue.",
            "P4: merchant descriptor inquiry without confirmed loss; standard review queue.",
        ],
    },
    "merchant_risk_standard_v1": {
        "title": "Merchant Risk Standard v1",
        "body": [
            "Merchant reference data must include a current status, category code, country and risk tier.",
            "Stale or conflicting merchant records require validation before a risk decision is made.",
        ],
    },
    "authentication_policy_v1": {
        "title": "Customer Authentication Policy v1",
        "body": [
            "Use step-up authentication when transaction context materially differs from recent customer behaviour.",
            "Repeated OTP or password failures must be available to fraud investigators.",
        ],
    },
    "investigation_procedure_v1": {
        "title": "Fraud Investigation Procedure v1",
        "body": [
            "Review the customer, account, card, payment, authentication, device and merchant records.",
            "Record the reasoning, evidence gaps, action taken and final disposition in investigator notes.",
        ],
    },
}


def _write_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    pdf.setTitle(title)
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(54, height - 64, title)
    y = height - 98
    pdf.setFont("Helvetica", 10.5)
    for paragraph in paragraphs:
        for line in wrap(paragraph, width=90):
            pdf.drawString(54, y, line)
            y -= 15
        y -= 8
    pdf.setFont("Helvetica-Oblique", 8)
    pdf.drawString(54, 40, "Synthetic document - Meridian Bank simulation - no real customer data")
    pdf.save()


def _write_image_only_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    """Create a PDF page made only of pixels, with no searchable text layer."""
    image = Image.new("RGB", (1240, 1754), "#f6f1e7")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=22)
    bold = ImageFont.load_default(size=31)
    draw.text((100, 100), title, fill="#1b2733", font=bold)
    y = 180
    for paragraph in paragraphs:
        for line in wrap(paragraph, width=78):
            draw.text((100, y), line, fill="#202832", font=font)
            y += 35
        y += 20
    draw.text((100, 1640), "Synthetic scanned policy copy", fill="#5d6570", font=font)
    temporary = path.with_suffix(".png")
    image.save(temporary)
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.drawImage(str(temporary), 0, 0, width=A4[0], height=A4[1])
    pdf.save()
    temporary.unlink()


def generate_documents(output_dir: Path) -> list[dict]:
    output_dir.mkdir(parents=True, exist_ok=True)
    metadata: list[dict] = []

    fraud = DOCUMENTS["fraud_policy_v1"]
    _write_pdf(output_dir / "fraud_policy_v1.pdf", fraud["title"], fraud["body"])
    metadata.append({"file": "fraud_policy_v1.pdf", "format": "PDF", "ocr_candidate": False})

    dispute = DOCUMENTS["dispute_procedure_v1"]
    docx = Document()
    docx.add_heading(dispute["title"], level=1)
    for paragraph in dispute["body"]:
        docx.add_paragraph(paragraph)
    docx.add_paragraph("Synthetic document - no real customer data")
    docx.save(output_dir / "dispute_procedure_v1.docx")
    metadata.append({"file": "dispute_procedure_v1.docx", "format": "DOCX", "ocr_candidate": False})

    escalation = DOCUMENTS["escalation_matrix_v1"]
    text = escalation["title"] + "\n\n" + "\n".join(escalation["body"]) + "\n"
    (output_dir / "escalation_matrix_v1.txt").write_text(text, encoding="utf-8")
    metadata.append({"file": "escalation_matrix_v1.txt", "format": "TXT", "ocr_candidate": False})

    merchant = DOCUMENTS["merchant_risk_standard_v1"]
    paragraphs = "\n".join(f"<p>{item}</p>" for item in merchant["body"])
    html = f"<!doctype html><html><head><meta charset='utf-8'><title>{merchant['title']}</title></head><body><h1>{merchant['title']}</h1>{paragraphs}<footer>Synthetic document</footer></body></html>"
    (output_dir / "merchant_risk_standard_v1.html").write_text(html, encoding="utf-8")
    metadata.append({"file": "merchant_risk_standard_v1.html", "format": "HTML", "ocr_candidate": False})

    auth = DOCUMENTS["authentication_policy_v1"]
    _write_pdf(output_dir / "authentication_policy_v1.pdf", auth["title"], auth["body"])
    metadata.append({"file": "authentication_policy_v1.pdf", "format": "PDF", "ocr_candidate": False})

    investigation = DOCUMENTS["investigation_procedure_v1"]
    _write_image_only_pdf(
        output_dir / "investigation_procedure_scanned_v1.pdf",
        investigation["title"],
        investigation["body"],
    )
    metadata.append(
        {
            "file": "investigation_procedure_scanned_v1.pdf",
            "format": "PDF_IMAGE_ONLY",
            "ocr_candidate": True,
        }
    )
    return metadata

